#!/usr/bin/env python3
"""
DOCX → Markdown 转换脚本
用法: python3 docx_to_md.py <docx_path>
输出: 同目录下同名 .md 文件
成功输出: CONVERTED:<md_path>
失败输出: ERROR:<原因> (stderr, exit 1)
"""
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR:python-docx 未安装，请运行: pip install python-docx", file=sys.stderr)
    sys.exit(1)

if len(sys.argv) < 2:
    print("ERROR:用法: python3 docx_to_md.py <docx_path>", file=sys.stderr)
    sys.exit(1)

docx_path = sys.argv[1]
md_path = str(Path(docx_path).with_suffix('.md'))

try:
    doc = Document(docx_path)
except Exception as e:
    print(f"ERROR:无法读取文件 {docx_path}: {e}", file=sys.stderr)
    sys.exit(1)

lines = []

for para in doc.paragraphs:
    style = para.style.name if para.style else ''
    text = para.text
    if not text.strip():
        lines.append('')
    elif 'Heading 1' in style:
        lines.append(f'# {text}')
    elif 'Heading 2' in style:
        lines.append(f'## {text}')
    elif 'Heading 3' in style:
        lines.append(f'### {text}')
    elif 'Heading 4' in style:
        lines.append(f'#### {text}')
    else:
        lines.append(text)

for table in doc.tables:
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        lines.append(' | '.join(cells))
        if i == 0:
            lines.append(' | '.join(['---'] * len(cells)))
    lines.append('')

with open(md_path, 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))

print(f'CONVERTED:{md_path}')
