"""
PRD Markdown → DOCX 转换器
支持：标题、段落、表格（含单元格内插图）、列表、粗体、代码、分隔线、Mermaid 流程图
支持：Mermaid 渲染方式选择（AI 高清 Seedream / 本地 Chrome headless）
"""
import re, json, sys, os, subprocess, tempfile, threading
import urllib.request
import base64

# ─── Seedream AI 图片生成 ───────────────────────────────────────────────

SEEDREAM_MODEL = 'doubao-seedream-4-5-251128'
SEEDREAM_URL = 'https://ark.cn-beijing.volces.com/api/v3/images/generations'
SEEDREAM_SIZE = '2560x1440'

STYLE_PRESETS = {
    'graph':           ('linear-progression', 'corporate-memphis', '蓝色系'),
    'sequenceDiagram': ('linear-progression', 'technical-schematic', '蓝色系'),
    'flowchart':       ('tree-branching', 'corporate-memphis', '蓝色系'),
    'classDiagram':    ('structural-breakdown', 'technical-schematic', '蓝色系'),
}

def _load_ark_api_key():
    """从 ~/.baoyu-skills/.env 或环境变量加载 ARK_API_KEY"""
    key = os.environ.get('ARK_API_KEY')
    if key:
        return key
    env_file = os.path.expanduser('~/.baoyu-skills/.env')
    if os.path.exists(env_file):
        for line in open(env_file):
            line = line.strip()
            if line.startswith('ARK_API_KEY='):
                return line.split('=', 1)[1].strip()
    return None

def generate_seedream(prompt, size=SEEDREAM_SIZE):
    """调用 Seedream API 生成图片，返回临时文件路径或 None"""
    api_key = _load_ark_api_key()
    if not api_key:
        print('  ⚠️ 未找到 ARK_API_KEY，无法使用 AI 生成')
        return None
    headers = {
        'Content-Type': 'application/json',
        'Authorization': f'Bearer {api_key}'
    }
    data = json.dumps({
        'model': SEEDREAM_MODEL,
        'prompt': prompt,
        'size': size,
        'response_format': 'b64_json'
    }).encode()
    try:
        req = urllib.request.Request(SEEDREAM_URL, data=data, headers=headers)
        resp = urllib.request.urlopen(req, timeout=120)
        result = json.loads(resp.read())
        img_data = base64.b64decode(result['data'][0]['b64_json'])
        out_png = tempfile.mktemp(suffix='.png', dir='/tmp')
        with open(out_png, 'wb') as f:
            f.write(img_data)
        return out_png
    except Exception as e:
        print(f'  ⚠️ Seedream 生成失败: {e}')
        return None

def _detect_mermaid_type(code):
    first_line = code.strip().split('\n')[0].strip()
    for key in STYLE_PRESETS:
        if first_line.startswith(key):
            return key
    return 'graph'

def _build_seedream_prompt(code, layout, style, palette):
    lines = [l.strip() for l in code.split('\n') if l.strip() and not l.strip().startswith('```')]
    flow_desc = '\n'.join(lines)
    return (
        f'专业产品流程信息图，扁平矢量{style}风格，纯白色背景(#FFFFFF)，{palette}配色(主色#1D4ED8)。'
        f'中文标注，清晰可读，简洁专业，适合嵌入PRD文档。充足留白，节点间用带箭头连接线。'
        f'布局类型：{layout}。'
        f'\n\n流程内容（基于以下 Mermaid 代码转化为可视化信息图）：\n{flow_desc}'
    )

def _prompt_mermaid_choice(code):
    """交互式选择 Mermaid 渲染方式。非 tty 时直接返回 local。"""
    if not sys.stdin.isatty():
        return 'local', None

    if not _load_ark_api_key():
        return 'local', None

    mtype = _detect_mermaid_type(code)
    layout, style, palette = STYLE_PRESETS.get(mtype, STYLE_PRESETS['graph'])

    print(f'\n  检测到流程图（{mtype}）')
    print(f'    A. AI 生成高清信息图（Seedream，会产生 API 费用）')
    print(f'    B. 本地渲染（Chrome headless，免费但质量一般）')
    choice = input('  选择 A/B（默认 B）：').strip().upper()
    if choice != 'A':
        return 'local', None

    ALT_STYLES = [
        ('tree-branching', 'corporate-memphis', '蓝色系'),
        ('hub-spoke', 'ikea-manual', '蓝灰色系'),
        ('linear-progression', 'technical-schematic', '蓝色系'),
    ]

    print(f'\n  推荐风格：{layout} × {style}（{palette}）')
    print(f'    1. 用推荐风格直接生成')
    print(f'    2. 换其他风格')
    print(f'    3. 自定义描述风格要求')
    sc = input('  选择（默认 1）：').strip()

    if sc == '2':
        for idx, (l, s, p) in enumerate(ALT_STYLES, 1):
            print(f'    {idx}. {l} × {s}（{p}）')
        ai = input('  选择编号（默认 1）：').strip()
        ai_idx = int(ai) - 1 if ai.isdigit() else 0
        if 0 <= ai_idx < len(ALT_STYLES):
            layout, style, palette = ALT_STYLES[ai_idx]
    elif sc == '3':
        custom = input('  描述你想要的风格：').strip()
        prompt = f'{custom}\n\n流程内容：\n{code}'
        return 'ai', prompt

    return 'ai', _build_seedream_prompt(code, layout, style, palette)

# ─── 原有代码 ───────────────────────────────────────────────────────────
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Chrome headless 二进制路径
CHROME = os.path.expanduser(
    '~/Library/Caches/ms-playwright/chromium-1212/chrome-mac-arm64/'
    'Google Chrome for Testing.app/Contents/MacOS/Google Chrome for Testing'
)

MERMAID_LOCAL = '/tmp/mermaid.min.js'   # 首次用时执行: curl -sL https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js -o /tmp/mermaid.min.js

# HTML 模板：内联 mermaid.js（避免 --allow-file-access-from-files 标志在 Python subprocess 下的兼容问题）
MERMAID_HTML_TMPL = '''<!DOCTYPE html><html><head><meta charset="UTF-8">
<style>* {{margin:0;padding:0;box-sizing:border-box;}}
body {{background:white;padding:16px;font-family:-apple-system,"PingFang SC",sans-serif;display:inline-block;}}</style>
</head><body><div class="mermaid">{code}</div>
<script>{mermaid_js_inline}</script>
<script>mermaid.initialize({{startOnLoad:true,theme:"default"}});</script>
</body></html>'''

def render_mermaid(code):
    """渲染 Mermaid 代码为 PNG（内联 JS，无需 --allow-file-access-from-files）"""
    if not os.path.exists(MERMAID_LOCAL):
        print(f'  ⚠️ 本地 mermaid.js 不存在，尝试下载...')
        try:
            subprocess.run(
                ['curl', '-sL', 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js',
                 '-o', MERMAID_LOCAL],
                timeout=30, check=True, capture_output=True
            )
        except Exception:
            return None

    # 读取 mermaid.js 并内联到 HTML（避免 file:// 跨源问题）
    mermaid_js_inline = open(MERMAID_LOCAL, encoding='utf-8').read()

    # 写临时 HTML 到 /tmp/
    html_file = tempfile.NamedTemporaryFile(suffix='.html', dir='/tmp', delete=False, mode='w', encoding='utf-8')
    html_file.write(MERMAID_HTML_TMPL.format(code=code, mermaid_js_inline=mermaid_js_inline))
    html_file.close()

    # Chrome headless 截图（内联 JS，无需 --allow-file-access-from-files）
    # shell=True 避免 macOS 下 Python subprocess list 模式与 Chrome 进程组的兼容问题
    # 渲染完成后等待 2s，让 Chrome 完全释放资源，避免多实例连续调用时 exit 133（SIGTRAP）
    out_png = tempfile.mktemp(suffix='.png', dir='/tmp')
    try:
        cmd = (f'"{CHROME}" --headless=new --no-sandbox --disable-gpu '
               f'--screenshot="{out_png}" --window-size=700,1200 '
               f'--hide-scrollbars --virtual-time-budget=8000 '
               f'"file://{html_file.name}" 2>/dev/null')
        subprocess.run(cmd, shell=True, timeout=30)
    except (subprocess.TimeoutExpired, Exception):
        return None
    finally:
        os.unlink(html_file.name)

    import time; time.sleep(2)

    if not os.path.exists(out_png):
        return None

    # 裁剪底部空白
    try:
        from PIL import Image, ImageChops
        img = Image.open(out_png).convert('RGB')
        bg = Image.new('RGB', img.size, (255, 255, 255))
        bbox = ImageChops.difference(img, bg).getbbox()
        if bbox:
            cropped = img.crop((0, 0, img.width, min(bbox[3] + 20, img.height)))
            cropped.save(out_png)
    except ImportError:
        pass  # Pillow 未安装时保留原图

    return out_png

def set_heading_style(para, level):
    run = para.runs[0] if para.runs else para.add_run()
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    colors = {1: (0x1d,0x1d,0x1f), 2: (0x2c,0x2c,0x2c), 3: (0x3c,0x3c,0x3c), 4: (0x4c,0x4c,0x4c)}
    run.font.size = Pt(sizes.get(level, 11))
    run.font.bold = True
    r, g, b = colors.get(level, (0,0,0))
    run.font.color.rgb = RGBColor(r, g, b)

def add_inline_formats(para, text):
    """处理粗体 **text**、行内代码 `code`、<br> 换行"""
    # 先按 <br> 切分成多段，段间插入真实换行符
    segments = re.split(r'<br\s*/?>', text, flags=re.IGNORECASE)
    for seg_idx, segment in enumerate(segments):
        if seg_idx > 0:
            para.add_run().add_break()
        parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', segment)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = para.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                if part:
                    para.add_run(part)

def shade_cell(cell, hex_color='F5F5F7'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def fill_cell(cell, text, is_header=False, screenshot_map=None):
    """
    填充单元格内容。
    若 text 包含 [xxx原型]，在单元格内插图（图在上，描述文字在下）。
    """
    cell.text = ''
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(2)

    # 检测原型占位符
    img_inserted = False
    if screenshot_map:
        m = re.search(r'\[(.+?)原型\]', text)
        if m:
            label = m.group(1)
            img_path = screenshot_map.get(label)
            if img_path:
                try:
                    run = para.add_run()
                    run.add_picture(img_path, width=Cm(13))
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_inserted = True
                    # 图片下方加描述文字（去掉占位符部分）
                    desc = re.sub(r'\[.+?原型\]\s*', '', text).strip()
                    if desc:
                        desc_para = cell.add_paragraph()
                        desc_run = desc_para.add_run(desc)
                        desc_run.font.size = Pt(9)
                        desc_run.font.color.rgb = RGBColor(0x86, 0x86, 0x8b)
                        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        desc_para.paragraph_format.space_before = Pt(4)
                except Exception as e:
                    print(f'  ⚠️ 图片插入失败 [{label}]: {e}')

    if not img_inserted:
        add_inline_formats(para, text)
        if is_header:
            for run in para.runs:
                run.bold = True
        for run in para.runs:
            run.font.size = Pt(10)

def set_col_width(cell, width_cm):
    """设置单元格固定宽度"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))  # 567 twips/cm
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def apply_col_widths(table, num_cols):
    """A4 正文宽约 16cm；2列时标题列4cm，内容列12cm；多列均分"""
    from docx.shared import Cm as _Cm
    PAGE_W = 16.0
    widths = [4.0, 12.0] if num_cols == 2 else [PAGE_W / num_cols] * num_cols
    # 设置 w:tblGrid（飞书优先读取此处）
    for i, col in enumerate(table.columns):
        if i < len(widths):
            col.width = _Cm(widths[i])
    # 同时设置每格 w:tcW（Office 读取此处）
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                set_col_width(cell, widths[i])

def add_table(doc, lines, screenshot_map=None):
    """解析 Markdown 表格并插入 DOCX，支持单元格内插图"""
    rows = [l for l in lines if not re.match(r'^\|[-| :]+\|$', l)]
    if len(rows) < 2:
        return
    def cells(row):
        return [c.strip() for c in row.split('|')[1:-1]]

    headers = cells(rows[0])
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # 表头行
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        fill_cell(cell, h, is_header=True)
        shade_cell(cell, 'F5F5F7')

    # 数据行
    for row_data in rows[1:]:
        row_cells = cells(row_data)
        row = table.add_row()
        for i, val in enumerate(row_cells):
            if i < len(row.cells):
                fill_cell(row.cells[i], val, screenshot_map=screenshot_map)

    # 设置列宽
    apply_col_widths(table, len(headers))
    doc.add_paragraph()

def _find_available_font(font_list):
    """从字体 fallback 链中找到第一个系统可用的字体"""
    for font in font_list:
        if font in ('serif', 'sans-serif', 'monospace', 'system'):
            return font  # CSS generic — 用 fallback 默认
        try:
            result = subprocess.run(['fc-list', font], capture_output=True, text=True, timeout=5)
            if result.stdout.strip():
                return font
        except Exception:
            pass
    return font_list[0] if font_list else 'PingFang SC'


def _parse_pt(val):
    """Parse '11pt' → 11, '16pt' → 16"""
    if isinstance(val, (int, float)):
        return int(val)
    m = re.match(r'(\d+(?:\.\d+)?)\s*pt', str(val))
    return int(float(m.group(1))) if m else 11


def convert(prd_path, output_path, manifest_path=None, recipe_config=None):
    prd_path = Path(prd_path)
    project_dir = prd_path.parent.parent

    # Recipe defaults
    body_font = 'PingFang SC'
    body_size = 11
    if recipe_config:
        body_font_chain = recipe_config.get('bodyFont', ['PingFang SC'])
        body_font = _find_available_font(body_font_chain)
        body_size = _parse_pt(recipe_config.get('bodySize', '11pt'))
        print(f'  配方字体: {body_font} ({body_size}pt)')

    # 加载截图映射 label → 本地路径
    screenshot_map = {}
    if manifest_path and Path(manifest_path).exists():
        manifest = json.loads(Path(manifest_path).read_text())
        for s in manifest['sections']:
            img_path = project_dir / '06-prototype' / s['screenshot']
            if img_path.exists():
                screenshot_map[s['label']] = str(img_path)
                print(f'  截图映射: [{s["label"]}原型] → {img_path.name}')

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = body_font
    style.font.size = Pt(body_size)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), body_font)

    md = prd_path.read_text(encoding='utf-8')
    lines = md.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # 标题
        m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            para = doc.add_heading('', level=min(level, 4))
            para.clear()
            add_inline_formats(para, m.group(2))
            set_heading_style(para, level)
            i += 1
            continue

        # 分隔线
        if re.match(r'^---+$', line.strip()):
            doc.add_paragraph('─' * 40)
            i += 1
            continue

        # Mermaid 流程图（支持 AI 高清 / 本地渲染 二选一）
        if re.match(r'^```mermaid\s*$', line.strip()):
            mermaid_lines = []
            i += 1
            while i < len(lines) and not re.match(r'^```\s*$', lines[i].strip()):
                mermaid_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束 ```
            code = '\n'.join(mermaid_lines)

            render_mode, ai_prompt = _prompt_mermaid_choice(code)

            if render_mode == 'ai' and ai_prompt:
                print(f'  🎨 AI 生成高清流程图中...')
                tmp_png = generate_seedream(ai_prompt)
                img_width = Cm(15)
                if tmp_png:
                    # 同时保存到 11-illustrations/
                    ill_dir = project_dir / '11-illustrations'
                    ill_dir.mkdir(exist_ok=True)
                    existing = sorted(ill_dir.glob('*.png'))
                    next_num = len(existing) + 1
                    save_name = f'{next_num:02d}-flowchart.png'
                    import shutil
                    shutil.copy2(tmp_png, str(ill_dir / save_name))
                    print(f'  ✅ AI 流程图已保存: 11-illustrations/{save_name}')
            else:
                print(f'  渲染 Mermaid 流程图（本地）...')
                tmp_png = render_mermaid(code)
                img_width = Cm(12)

            if tmp_png:
                para = doc.add_paragraph()
                run = para.add_run()
                run.add_picture(tmp_png, width=img_width)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                os.unlink(tmp_png)
            else:
                para = doc.add_paragraph(code)
                para.style = 'No Spacing'
            continue

        # 表格
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, table_lines, screenshot_map=screenshot_map)
            continue

        # 无序列表
        if re.match(r'^[-*]\s+', line):
            para = doc.add_paragraph(style='List Bullet')
            para.clear()
            add_inline_formats(para, re.sub(r'^[-*]\s+', '', line))
            i += 1
            continue

        # 有序列表
        if re.match(r'^\d+\.\s+', line):
            para = doc.add_paragraph(style='List Number')
            para.clear()
            add_inline_formats(para, re.sub(r'^\d+\.\s+', '', line))
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # Markdown 图片语法 ![alt](path)
        m_img = re.match(r'^!\[(.+?)\]\((.+?)\)$', line.strip())
        if m_img:
            alt_text = m_img.group(1)
            img_rel = m_img.group(2)
            # 相对于 PRD 文件所在目录解析路径
            img_abs = (prd_path.parent / img_rel).resolve()
            if img_abs.exists():
                try:
                    para = doc.add_paragraph()
                    run = para.add_run()
                    run.add_picture(str(img_abs), width=Cm(15))
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption = doc.add_paragraph(alt_text)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in caption.runs:
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(0x86, 0x86, 0x8b)
                    print(f'  ✅ 图片已嵌入: {alt_text} ({img_abs.name})')
                except Exception as e:
                    print(f'  ⚠️ 图片插入失败 [{alt_text}]: {e}')
            else:
                print(f'  ⚠️ 图片不存在: {img_abs}')
                para = doc.add_paragraph()
                add_inline_formats(para, f'[图片: {alt_text}]')
            i += 1
            continue

        # 原型截图占位符（独立行 [xxx原型]）
        m_proto = re.match(r'^\[(.+?)原型\]$', line.strip())
        if m_proto and screenshot_map:
            label = m_proto.group(1)
            img_path = screenshot_map.get(label)
            if img_path:
                try:
                    para = doc.add_paragraph()
                    run = para.add_run()
                    run.add_picture(img_path, width=Cm(14))
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption = doc.add_paragraph(label + ' 原型截图')
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in caption.runs:
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(0x86, 0x86, 0x8b)
                    print(f'  ✅ 截图已嵌入: [{label}]')
                except Exception as e:
                    print(f'  ⚠️ 截图插入失败 [{label}]: {e}')
                i += 1
                continue

        # 普通段落
        para = doc.add_paragraph()
        add_inline_formats(para, line)
        para.paragraph_format.space_after = Pt(6)
        i += 1

    doc.save(output_path)
    print(f'✅ DOCX 已生成: {output_path}')

if __name__ == '__main__':
    # 用法: python3 md2docx.py <prd.md> <output.docx> [manifest.json] [--recipe=NAME] [--recipe-file=PATH]
    if len(sys.argv) < 3:
        print('用法: python3 md2docx.py <prd.md> <output.docx> [manifest.json] [--recipe=NAME] [--recipe-file=PATH]')
        sys.exit(1)

    positional = [a for a in sys.argv[1:] if not a.startswith('--')]
    named = {a.split('=', 1)[0]: a.split('=', 1)[1] for a in sys.argv[1:] if '=' in a and a.startswith('--')}

    prd_file = positional[0]
    out_file = positional[1]
    manifest_file = positional[2] if len(positional) > 2 else None

    recipe_cfg = None
    recipe_name = named.get('--recipe')
    recipe_file = named.get('--recipe-file')
    if recipe_name and recipe_file and Path(recipe_file).exists():
        try:
            all_recipes = json.loads(Path(recipe_file).read_text())
            recipe_cfg = all_recipes.get('recipes', {}).get(recipe_name)
            if recipe_cfg:
                print(f'  使用配方: {recipe_name} ({recipe_cfg.get("label", "")})')
            else:
                print(f'  ⚠️ 配方 "{recipe_name}" 不存在，使用默认样式')
        except Exception as e:
            print(f'  ⚠️ 读取配方文件失败: {e}，使用默认样式')

    convert(prd_file, out_file, manifest_file, recipe_config=recipe_cfg)
